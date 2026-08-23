"""Safe, deliberately small document parsers for the supported v1 formats."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from rag_service.domain import ParsedDocument
from rag_service.errors import UnsupportedDocumentError, ValidationError

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_MARKDOWN_MIMES = {"text/markdown", "text/x-markdown"}
_TEXT_MIMES = {"text/plain", *_MARKDOWN_MIMES}
_EXTENSION_MIMES = {
    ".pdf": "application/pdf",
    ".docx": _DOCX_MIME,
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".markdown": "text/markdown",
}


def normalized_content_type(filename: str, supplied: str | None) -> str:
    supplied = (supplied or "").lower().split(";", maxsplit=1)[0].strip()
    inferred = _EXTENSION_MIMES.get(Path(filename).suffix.lower())
    if supplied in {"", "application/octet-stream"} and inferred:
        return inferred
    if supplied == "application/pdf" or supplied == _DOCX_MIME or supplied in _TEXT_MIMES:
        return supplied
    if inferred and supplied == inferred:
        return supplied
    raise UnsupportedDocumentError(f"Unsupported document type for {filename}")


def parse_document(filename: str, content_type: str, content: bytes) -> ParsedDocument:
    if content_type == "application/pdf":
        if not content.startswith(b"%PDF"):
            raise ValidationError(f"{filename} is not a valid PDF")
        try:
            reader = PdfReader(BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages]
        except Exception as error:
            raise ValidationError(f"{filename} is not a valid PDF") from error
        text = "\n\n".join(f"[Page {index + 1}]\n{page}" for index, page in enumerate(pages))
        if not text.strip() or not any(page.strip() for page in pages):
            raise UnsupportedDocumentError(f"{filename} appears to be a scanned PDF; OCR is not enabled")
        return ParsedDocument(text=text, page_count=len(pages))
    if content_type == _DOCX_MIME:
        if not content.startswith(b"PK"):
            raise ValidationError(f"{filename} is not a valid DOCX file")
        document = DocxDocument(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        if not text.strip():
            raise ValidationError(f"{filename} has no extractable text")
        return ParsedDocument(text=text)
    if content_type in _TEXT_MIMES:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError as error:
            raise ValidationError(f"{filename} must be UTF-8 encoded") from error
        if not text.strip():
            raise ValidationError(f"{filename} has no text")
        return ParsedDocument(text=text)
    raise UnsupportedDocumentError(f"Unsupported document type for {filename}")
